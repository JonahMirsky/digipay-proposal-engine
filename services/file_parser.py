"""Extract text from uploaded files."""

import csv
import io

from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook

ALLOWED = {'.pdf', '.docx', '.xlsx', '.csv', '.txt'}


def parse_file(file_storage):
    filename = file_storage.filename or 'unknown'
    ext = ('.' + filename.rsplit('.', 1)[-1]).lower() if '.' in filename else ''

    if ext not in ALLOWED:
        raise ValueError(f'Unsupported file type: {ext}')

    data = file_storage.read()
    extractors = {
        '.pdf': _pdf, '.docx': _docx, '.xlsx': _xlsx, '.csv': _csv, '.txt': _txt,
    }

    text = extractors[ext](data)
    if not text.strip():
        raise ValueError('No text could be extracted from this file.')

    return {'text': text.strip(), 'filename': filename}


def _pdf(data):
    pages = []
    for page in PdfReader(io.BytesIO(data)).pages:
        t = page.extract_text()
        if t:
            pages.append(t)
    return '\n\n'.join(pages)


def _docx(data):
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(c.text.strip() for c in row.cells))
    return '\n'.join(parts)


def _xlsx(data):
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for name in wb.sheetnames:
        lines.append(f'--- {name} ---')
        for row in wb[name].iter_rows(values_only=True):
            cells = [str(c) if c is not None else '' for c in row]
            if any(cells):
                lines.append(' | '.join(cells))
    return '\n'.join(lines)


def _csv(data):
    lines = []
    for row in csv.reader(io.StringIO(data.decode('utf-8', errors='replace'))):
        lines.append(' | '.join(row))
    return '\n'.join(lines)


def _txt(data):
    return data.decode('utf-8', errors='replace')
